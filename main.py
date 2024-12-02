import streamlit as st
import json
import redis
from utils.pdf_processing import process_pdf_task
from utils.respondent import ask_question
from utils.config import redis_host, redis_pass
import io
from docx import Document
import uuid
import tiktoken
from docx.shared import Pt
import re

def remove_markdown(text):
    """Remove Markdown formatting from text."""
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text

def count_tokens(text, model="gpt-4o"):
    encoding = tiktoken.encoding_for_model(model)
    tokens = encoding.encode(text)
    return len(tokens)

# Initialize Redis client
redis_client = redis.Redis(
    host=redis_host,
    port=6379,
    password=redis_pass,
)

# Initialize session state
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "doc_token" not in st.session_state:
    st.session_state.doc_token = 0

def save_document_to_redis(session_id, file_name, document_data):
    redis_key = f"{session_id}:document_data:{file_name}"
    redis_client.set(redis_key, json.dumps(document_data))

def get_document_from_redis(session_id, file_name):
    redis_key = f"{session_id}:document_data:{file_name}"
    data = redis_client.get(redis_key)
    if data:
        return json.loads(data)
    return None

def retrieve_user_documents_from_redis(session_id):
    documents = {}
    for key in redis_client.keys(f"{session_id}:document_data:*"):
        file_name = key.decode().split(f"{session_id}:document_data:")[1]
        documents[file_name] = get_document_from_redis(session_id, file_name)
    return documents

def handle_question(prompt, spinner_placeholder):
    if prompt:
        try:
            documents_data = retrieve_user_documents_from_redis(
                st.session_state.session_id
            )
            with spinner_placeholder.container():
                st.spinner('Thinking...')
                answer, tot_tokens = ask_question(
                    documents_data, prompt, st.session_state.chat_history
                )
            st.session_state.chat_history.append(
                {
                    "question": prompt,
                    "answer": f"{answer}\nTotal tokens: {tot_tokens}",
                }
            )
            #st.success("Here is your answer!")
        except Exception as e:
            st.error(f"Error processing question: {e}")
        finally:
            spinner_placeholder.empty()

def reset_session():
    st.session_state.chat_history = []
    st.session_state.doc_token = 0
    for key in redis_client.keys(f"{st.session_state.session_id}:document_data:*"):
        redis_client.delete(key)

def display_chat():
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history):
            with st.chat_message("user"):
                st.write(chat['question'])
            with st.chat_message("assistant"):
                st.write(chat['answer'])
            chat_content = {
                "question": chat["question"],
                "answer": chat["answer"],
            }
            doc = generate_word_document(chat_content)
            word_io = io.BytesIO()
            doc.save(word_io)
            word_io.seek(0)
            st.download_button(
                label="Download",
                data=word_io,
                file_name=f"chat_{i+1}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

def apply_markdown_to_word(doc, markdown_text):
    lines = markdown_text.splitlines()
    for line in lines:
        heading_match = re.match(r"^(#{1,4})\s*(.+)", line)
        if heading_match:
            header_level = len(heading_match.group(1))
            header_text = heading_match.group(2)

            header_para = doc.add_paragraph()
            header_run = header_para.add_run(header_text)
            header_run.bold = True
            if header_level == 1:
                header_run.font.size = Pt(14)
            elif header_level == 2:
                header_run.font.size = Pt(14)
            elif header_level == 3:
                header_run.font.size = Pt(11)
            else:
                header_run.font.size = Pt(14)
            continue

        if not line.strip():
            continue

        paragraph = doc.add_paragraph()

        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", line)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            else:
                paragraph.add_run(part)

def generate_word_document(content):
    doc = Document()

    heading = doc.add_heading("Chat Response", level=0)
    heading.runs[0].font.name = "Aptos"
    heading.runs[0].font.size = Pt(14)

    apply_markdown_to_word(doc, f"Question: {content['question']}")
    apply_markdown_to_word(doc, f"Answer: {content['answer']}")

    return doc

# Main UI
st.image("logoD.png", width=200)
st.title("docQuest")
st.subheader("Unveil the Essence, Compare Easily, Analyze Smartly")

with st.sidebar:
    uploaded_files = st.file_uploader(
        "Upload files less than 400 pages",
        type=["pdf", "docx", "xlsx", "pptx"],
        accept_multiple_files=True,
        help="If your question is not answered properly or there's an error, consider uploading smaller documents or splitting larger ones.",
    )

    if uploaded_files:
        new_files = []
        for uploaded_file in uploaded_files:
            if not redis_client.exists(
                f"{st.session_state.session_id}:document_data:{uploaded_file.name}"
            ):
                new_files.append(uploaded_file)
            else:
                st.info(f"{uploaded_file.name} is already processed.")

        if new_files:
            progress_text = st.empty()
            progress_bar = st.progress(0)
            total_files = len(new_files)

            with st.spinner("Learning about your document(s)..."):
                try:
                    for i, uploaded_file in enumerate(new_files):
                        document_data = process_pdf_task(uploaded_file, first_file=(i == 0))
                        st.session_state.doc_token += count_tokens(str(document_data))
                        if st.session_state.doc_token>4000:
                            st.warning('Document is too large to query, results may be inaccurate. Consider uploading smaller document.', icon="⚠️")
                        save_document_to_redis(
                            st.session_state.session_id,
                            uploaded_file.name,
                            document_data,
                        )
                        st.success(f"{uploaded_file.name} processed!")
                        progress_bar.progress((i + 1) / total_files)
                except Exception as e:
                    st.error(f"Error processing file: {e}")

            progress_text.text("Processing complete.")
            progress_bar.empty()

    if retrieve_user_documents_from_redis(st.session_state.session_id):
        download_data = json.dumps(
            retrieve_user_documents_from_redis(st.session_state.session_id), indent=4
        )
        st.download_button(
            label="Download Document Analysis",
            data=download_data,
            file_name="document_analysis.json",
            mime="application/json",
        )

if retrieve_user_documents_from_redis(st.session_state.session_id):
    # Use st.chat_input for user input
    prompt = st.chat_input("Ask me anything about your documents", key="chat_input")
    spinner_placeholder = st.empty()
    if prompt:
        handle_question(prompt, spinner_placeholder)

display_chat()
